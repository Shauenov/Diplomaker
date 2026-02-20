
try:
    from docx import Document
    print("Loading docx...")
    doc = Document(r'c:\Users\user\OneDrive\Рабочий стол\template\рус.docx')
    print(f"Loaded. Paragraphs: {len(doc.paragraphs)}")
    
    with open('docx_dump.txt', 'w', encoding='utf-8') as f:
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                f.write(text + "\n")
                print(f"Extracted: {text[:50]}...")
        
        # Also check tables
        print(f"Tables: {len(doc.tables)}")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        f.write(text + "\n")
                        print(f"Table Cell: {text[:50]}...")
                        
except Exception as e:
    print(f"Error: {e}")
